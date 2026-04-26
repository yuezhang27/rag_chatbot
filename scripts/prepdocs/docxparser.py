"""
DOCX text extraction. Supports local (python-docx) and Azure Document Intelligence.

backend="local": python-docx 段落 + 表格提取（开发用）
backend="azure": Azure Document Intelligence Layout 模型（生产用）
    - 复用 pdfparser.py 的 Azure DI 调用（DI 原生支持 DOCX）
"""
import io
import logging
from typing import List, Literal, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Local parser (python-docx)
# ---------------------------------------------------------------------------


def _docx_table_to_markdown(table) -> str:
    """将 python-docx Table 对象转为 Markdown 表格字符串。

    输出格式与 pdfparser._table_to_markdown() 一致。
    """
    rows = table.rows
    if not rows:
        return ""

    lines: List[str] = []
    for i, row in enumerate(rows):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("| " + " | ".join("---" for _ in cells) + " |")

    return "\n".join(lines)


def parse_docx_local_pages(
    content: bytes, filename: str = ""
) -> List[Tuple[int, str]]:
    """使用 python-docx 提取段落和表格，按 Heading 分段。

    分段策略：遇到 Heading 1 或 Heading 2 时开始新 section。
    若文档无 Heading，则整篇作为 section 1。

    返回值：[(section_number, section_text), ...]
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for local DOCX parsing. "
            "Install it with: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        logger.warning("Failed to open DOCX %s: %s", filename, exc)
        return []

    # Collect all block-level elements (paragraphs and tables) in document order.
    # python-docx exposes doc.element.body which contains both.
    sections: List[Tuple[int, str]] = []
    current_parts: List[str] = []
    section_num = 1

    # Build a set of table elements for quick lookup
    table_elements = {tbl._element for tbl in doc.tables}
    table_map = {tbl._element: tbl for tbl in doc.tables}

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # It's a paragraph
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            style_name = (para.style.name or "") if para.style else ""

            # Start new section on Heading 1 or Heading 2
            if style_name in ("Heading 1", "Heading 2"):
                if current_parts:
                    text = "\n\n".join(current_parts).strip()
                    if text:
                        sections.append((section_num, text))
                    section_num += 1
                    current_parts = []

            text = para.text.strip()
            if text:
                current_parts.append(text)

        elif tag == "tbl" and child in table_map:
            # It's a table
            md = _docx_table_to_markdown(table_map[child])
            if md:
                current_parts.append(md)

    # Flush remaining content
    if current_parts:
        text = "\n\n".join(current_parts).strip()
        if text:
            sections.append((section_num, text))

    if not sections:
        logger.warning("DOCX %s has no extractable content", filename)

    return sections


# ---------------------------------------------------------------------------
# Azure Document Intelligence parser
# ---------------------------------------------------------------------------


def parse_docx_azure_pages(
    content: bytes, filename: str = ""
) -> List[Tuple[int, str]]:
    """使用 Azure Document Intelligence 解析 DOCX（DI 原生支持 DOCX 格式）。

    复用 pdfparser.py 中的 DI 调用和 _build_page_content_with_tables()。
    """
    from scripts.prepdocs.pdfparser import (
        _get_di_credentials,
        _build_page_content_with_tables,
        _DI_POLLING_TIMEOUT,
    )

    endpoint, key = _get_di_credentials()

    from azure.ai.documentintelligence import DocumentIntelligenceClient
    from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
    from azure.core.credentials import AzureKeyCredential

    client = DocumentIntelligenceClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )
    poller = client.begin_analyze_document(
        model_id="prebuilt-layout",
        analyze_request=AnalyzeDocumentRequest(bytes_source=content),
    )
    result = poller.result(timeout=_DI_POLLING_TIMEOUT)

    return _build_page_content_with_tables(result)


# ---------------------------------------------------------------------------
# Public dispatch function
# ---------------------------------------------------------------------------


def parse_docx_pages(
    content: bytes,
    filename: str = "",
    backend: Literal["local", "azure"] = "local",
) -> List[Tuple[int, str]]:
    """按节/页解析 DOCX，供 ingestion pipeline 使用。

    backend="local": python-docx 提取，返回 (section_number, text)
    backend="azure": Azure DI 提取，返回 (page_number, text)
    """
    if backend == "local":
        return parse_docx_local_pages(content, filename)
    if backend == "azure":
        return parse_docx_azure_pages(content, filename)
    raise ValueError("backend must be 'local' or 'azure'")
